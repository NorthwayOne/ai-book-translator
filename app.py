"""
PDF Translator App
==================
Загрузи PDF → извлеки текст → переводи chunk за chunk → сохрани в JSON или DOCX.

Установка зависимостей:
    pip install streamlit pymupdf python-docx pyperclip

Запуск:
    streamlit run app.py
"""

import streamlit as st
import pymupdf          # pip install pymupdf
import json
import re
from pathlib import Path
from docx import Document   # pip install python-docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Файл проекта — лежит рядом с app.py
PROJECT_FILE = Path(__file__).parent / "project.json"


# ──────────────────────────────────────────────
#  УТИЛИТЫ
# ──────────────────────────────────────────────

def _is_corrupted_page(text: str) -> bool:
    """
    Проверяет, является ли текст страницы «мусорным».
    Считает долю нормальных символов (буквы, цифры, пунктуация).
    Если доля < 60% или текст слишком короткий — страница помечается как corrupted.
    """
    if len(text.strip()) < 30:          # почти пустая страница
        return True
    normal = sum(1 for c in text if c.isalnum() or c in ' .,;:!?-–—\'"()\n\t')
    ratio  = normal / max(len(text), 1)
    return ratio < 0.60


def extract_text_from_pdf(uploaded_file) -> tuple[str, list[int]]:
    """
    Читает PDF из загруженного файла.
    Возвращает (полный_текст, список_номеров_проблемных_страниц).

    Для каждой страницы:
      1. Извлекаем текст через pymupdf.
      2. Проверяем качество через _is_corrupted_page().
      3. Если плохо → вставляем placeholder [UNREADABLE TEXT — стр. N].
         Страница логируется в corrupted_pages.
    """
    pdf_bytes = uploaded_file.read()
    doc       = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    pages_text      = []
    corrupted_pages = []

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        if _is_corrupted_page(text):
            corrupted_pages.append(page_num)
            pages_text.append(f"[UNREADABLE TEXT — стр. {page_num}]")
        else:
            pages_text.append(text)

    doc.close()
    return "\n".join(pages_text), corrupted_pages


def split_into_chunks(text: str, words_per_chunk: int = 1000, overlap: int = 150) -> list[str]:
    """Разбивает текст на части по ~words_per_chunk слов с перекрытием overlap слов."""
    words = text.split()
    chunks = []
    step = words_per_chunk - overlap   # шаг сдвига с учётом перекрытия
    for i in range(0, len(words), step):
        chunk = " ".join(words[i : i + words_per_chunk])
        chunks.append(chunk)
        if i + words_per_chunk >= len(words):
            break
    return chunks


# Паттерн для поиска заголовков глав.
# Поддерживает: CHAPTER, Chapter, ГЛАВА — в начале строки,
# за которым следует номер или слово (1, One, I, ПЕРВАЯ и т.д.)
_CHAPTER_RE = re.compile(
    r'(?:^|\n)[ \t]*'                         # начало строки
    r'((?:CHAPTER|Chapter|ГЛАВА)'             # ключевое слово
    r'[ \t]+\S[^\n]{0,80})'                   # номер + опциональный подзаголовок
    r'[ \t]*(?:\n|$)',                         # конец строки
)


def detect_chapters(chunks: list[str]) -> list[dict]:
    """
    Сканирует список чанков и находит заголовки глав.
    Возвращает список:
        [{"chunk_idx": int, "chapter_num": int, "title": str}, ...]

    Дубликаты из overlap-зон отфильтровываются по тексту заголовка.
    """
    chapters   = []
    chapter_n  = 0
    seen       = set()   # нормализованные заголовки уже найденных глав

    for chunk_idx, text in enumerate(chunks):
        match = _CHAPTER_RE.search(text)
        if not match:
            continue
        title     = match.group(1).strip()
        title_key = title.lower()           # регистронезависимое сравнение
        if title_key in seen:
            continue                        # дубликат из перекрытия — пропускаем
        seen.add(title_key)
        chapter_n += 1
        chapters.append({
            "chunk_idx":  chunk_idx,
            "chapter_num": chapter_n,
            "title":       title,
        })

    return chapters


def chunk_to_chapter(chunk_idx: int, chapters: list[dict]) -> dict | None:
    """
    Возвращает главу, которой принадлежит chunk_idx.
    Логика: последняя глава, чей chunk_idx ≤ текущему.
    Возвращает None, если глав нет или чанк до первой главы.
    """
    result = None
    for ch in chapters:
        if ch["chunk_idx"] <= chunk_idx:
            result = ch
        else:
            break
    return result


def parse_glossary(raw: str) -> dict[str, str]:
    """
    Разбирает глоссарий из текста вида:
        English term → Russian translation
    Возвращает словарь {английский_термин: русский_перевод}.
    Строки без '→' игнорируются.
    """
    glossary = {}
    for line in raw.splitlines():
        if "→" in line:
            en, _, ru = line.partition("→")
            en, ru = en.strip(), ru.strip()
            if en and ru:
                glossary[en] = ru
    return glossary


def build_prompt(chunk_text: str, chunk_number: int, glossary: dict[str, str]) -> str:
    """
    Формирует промпт для перевода фрагмента научно-популярного текста.
    Если глоссарий не пустой — добавляет раздел с терминами.
    """

    # ── Раздел глоссария (подставляется автоматически) ──
    if glossary:
        glossary_lines = "\n".join(f"  {en} → {ru}" for en, ru in glossary.items())
        glossary_block = (
            "ГЛОССАРИЙ ТЕРМИНОВ\n"
            "Используй эти переводы последовательно по всему тексту:\n"
            f"{glossary_lines}\n\n"
        )
    else:
        glossary_block = ""

    return (
        f"Ты — профессиональный переводчик научно-популярной литературы.\n"
        f"Переведи фрагмент {chunk_number} с английского на русский язык.\n\n"
        f"ПРАВИЛА ПЕРЕВОДА:\n"
        f"1. Сохраняй точный смысл: не добавляй и не убирай информацию.\n"
        f"2. Сохраняй структуру: абзацы, отступы, списки, заголовки — без изменений.\n"
        f"3. Единообразие терминов: каждый термин переводи одинаково во всём тексте.\n"
        f"4. Не сокращай текст: каждое предложение оригинала должно быть в переводе.\n"
        f"5. Язык перевода — грамотный литературный русский, без канцеляризмов.\n"
        f"6. Выведи ТОЛЬКО перевод — без предисловий, пояснений и комментариев.\n\n"
        f"{glossary_block}"
        f"ТЕКСТ ДЛЯ ПЕРЕВОДА:\n{chunk_text}"
    )


def save_translations_to_json(chunks: list[str], translations: dict) -> str:
    """Сериализует chunks и переводы в JSON-строку."""
    data = []
    for i, chunk in enumerate(chunks):
        data.append({
            "chunk_number": i + 1,
            "original":     chunk,
            "translation":  translations.get(i, ""),
        })
    return json.dumps(data, ensure_ascii=False, indent=2)


def assemble_book(
    chunks: list[str],
    translations: dict,
    chapters: list[dict],
) -> list[dict]:
    """
    Склеивает чанки обратно в структуру книги по главам.

    Возвращает список разделов:
        [{"title": str, "ru_text": str, "en_text": str}, ...]

    Чанки до первой главы объединяются в раздел «Предисловие».
    Чанки между главами объединяются в соответствующую главу.
    """
    if not chapters:
        # Нет глав — вся книга как один раздел
        ru = "\n\n".join(translations.get(i, "") for i in range(len(chunks))).strip()
        en = "\n\n".join(chunks).strip()
        return [{"title": "Текст", "ru_text": ru, "en_text": en}]

    # Строим границы глав: chapter[i] занимает чанки [start_idx, next_start - 1]
    sections = []
    for i, ch in enumerate(chapters):
        start = ch["chunk_idx"]
        end   = chapters[i + 1]["chunk_idx"] if i + 1 < len(chapters) else len(chunks)
        title = f"Глава {ch['chapter_num']}. {ch['title']}"

        ru_parts = [translations.get(j, "") for j in range(start, end)]
        en_parts = [chunks[j]               for j in range(start, end)]

        sections.append({
            "title":   title,
            "ru_text": "\n\n".join(p for p in ru_parts if p).strip(),
            "en_text": "\n\n".join(en_parts).strip(),
        })

    # Чанки до первой главы → Предисловие
    first_ch_idx = chapters[0]["chunk_idx"]
    if first_ch_idx > 0:
        ru_pre = "\n\n".join(translations.get(j, "") for j in range(first_ch_idx)).strip()
        en_pre = "\n\n".join(chunks[j]               for j in range(first_ch_idx)).strip()
        if ru_pre or en_pre:
            sections.insert(0, {
                "title":   "Предисловие",
                "ru_text": ru_pre,
                "en_text": en_pre,
            })

    return sections


# Паттерны для автоматического выделения в тексте
_SPECIAL_PATTERNS = re.compile(
    r'^(IMPORTANT:|NOTE:|WARNING:|ВНИМАНИЕ:|ПРИМЕЧАНИЕ:)',
    re.MULTILINE,
)


def _setup_book_styles(doc: Document) -> None:
    """
    Настраивает базовые стили документа для книжного вида.
    Работает с существующими стилями python-docx без создания новых.
    """
    # Стиль Normal — основной текст книги
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    nfmt = normal.paragraph_format
    nfmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    nfmt.first_line_indent = Cm(1.25)
    nfmt.space_after       = Pt(0)
    nfmt.line_spacing      = Pt(18)

    # Heading 1 — названия глав
    h1 = doc.styles["Heading 1"]
    h1.font.name      = "Times New Roman"
    h1.font.size      = Pt(18)
    h1.font.bold      = True
    h1.font.color.rgb = None    # чёрный
    h1fmt = h1.paragraph_format
    h1fmt.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    h1fmt.space_before = Pt(24)
    h1fmt.space_after  = Pt(18)

    # Heading 2 — подзаголовки разделов (РУССКАЯ КНИГА / ОРИГИНАЛ)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2fmt = h2.paragraph_format
    h2fmt.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    h2fmt.space_before = Pt(36)
    h2fmt.space_after  = Pt(24)


def _add_book_paragraph(doc: Document, text: str) -> None:
    """
    Добавляет абзац с книжным форматированием.
    Строки с IMPORTANT: / NOTE: / WARNING: выделяются жирным курсивом.
    Пустые строки пропускаются (контролируем отступы через стиль).
    """
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        para = doc.add_paragraph(style="Normal")

        # Проверяем на специальные паттерны
        m = _SPECIAL_PATTERNS.match(stripped)
        if m:
            keyword = m.group(1)
            rest    = stripped[len(keyword):].strip()
            # Ключевое слово — жирный курсив
            kw_run = para.add_run(keyword + " ")
            kw_run.bold   = True
            kw_run.italic = True
            # Остальной текст — обычный, но курсив
            if rest:
                rest_run        = para.add_run(rest)
                rest_run.italic = True
        else:
            para.add_run(stripped)


def export_to_docx(
    chunks: list[str],
    translations: dict,
    chapters: list[dict] | None = None,
    book_title: str = "Перевод книги",
) -> BytesIO:
    """
    Экспортирует переведённую книгу в профессиональный DOCX.

    Структура документа:
      ── Титульная страница
      ── РУССКАЯ КНИГА
         Глава 1 / Предисловие / ...
         текст подряд (без chunk-разделителей)
      ── ОРИГИНАЛ НА АНГЛИЙСКОМ
         Chapter 1 / ...
         английский текст подряд
    """
    if chapters is None:
        chapters = []

    doc = Document()
    _setup_book_styles(doc)

    # Поля страницы (2.5 cm со всех сторон — книжный стандарт)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # ── Титульная страница ───────────────────────
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    title_run = title_para.add_run(book_title)
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(24)
    title_run.bold      = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Перевод с английского языка")
    sub_run.font.name = "Times New Roman"
    sub_run.font.size = Pt(14)
    sub_run.italic    = True

    doc.add_page_break()

    # Собираем книгу из чанков
    book_sections = assemble_book(chunks, translations, chapters)

    # ── РУССКАЯ КНИГА ────────────────────────────
    doc.add_heading("РУССКАЯ КНИГА", level=2)
    doc.add_page_break()

    for section in book_sections:
        if not section["ru_text"]:
            continue
        doc.add_heading(section["title"], level=1)
        _add_book_paragraph(doc, section["ru_text"])
        doc.add_page_break()

    # ── ОРИГИНАЛ НА АНГЛИЙСКОМ ───────────────────
    doc.add_heading("ОРИГИНАЛ НА АНГЛИЙСКОМ", level=2)
    doc.add_page_break()

    for section in book_sections:
        if not section["en_text"]:
            continue
        doc.add_heading(section["title"], level=1)
        _add_book_paragraph(doc, section["en_text"])
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ──────────────────────────────────────────────
#  PERSISTENT STORAGE
# ──────────────────────────────────────────────

def save_project() -> None:
    """
    Сохраняет весь проект в project.json рядом с app.py.
    Вызывается автоматически после каждого сохранения перевода.
    """
    # translations хранит ключи-int; JSON требует строки — конвертируем
    translations_str = {
        str(k): v for k, v in st.session_state.translations.items()
    }
    data = {
        "pdf_name":         st.session_state.pdf_name,
        "chunks":           st.session_state.chunks,
        "translations":     translations_str,
        "glossary":         st.session_state.glossary,
        "chapters":         st.session_state.chapters,
        "corrupted_pages":  st.session_state.corrupted_pages,
        "completed_chunks": sorted(st.session_state.completed_chunks),
        "current_chunk":    st.session_state.active_chunk,
    }
    PROJECT_FILE.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def load_project() -> bool:
    """
    Загружает проект из project.json в session_state.
    Возвращает True, если файл найден и успешно прочитан.
    """
    if not PROJECT_FILE.exists():
        return False
    try:
        data = json.loads(PROJECT_FILE.read_text(encoding="utf-8"))
        st.session_state.pdf_name         = data.get("pdf_name", "")
        st.session_state.chunks           = data.get("chunks", [])
        st.session_state.active_chunk     = data.get("current_chunk", 0)
        st.session_state.completed_chunks = set(data.get("completed_chunks", []))
        st.session_state.chapters         = data.get("chapters", [])
        st.session_state.corrupted_pages  = data.get("corrupted_pages", [])
        # Ключи приходят строками из JSON — переводим обратно в int
        st.session_state.translations = {
            int(k): v for k, v in data.get("translations", {}).items()
        }
        # Поддержка нового формата (dict) и старого (glossary_raw строка)
        if "glossary" in data:
            st.session_state.glossary = data["glossary"]
        elif "glossary_raw" in data:
            st.session_state.glossary = parse_glossary(data["glossary_raw"])
        else:
            st.session_state.glossary = {}
        return True
    except Exception:
        return False          # битый файл — стартуем с чистого листа



# ──────────────────────────────────────────────
#  ИНИЦИАЛИЗАЦИЯ SESSION STATE
# ──────────────────────────────────────────────

def init_state():
    defaults = {
        "chunks":            [],      # список строк-фрагментов
        "translations":      {},      # {индекс_чанка: строка перевода}
        "active_chunk":      0,       # выбранный чанк
        "pdf_name":          "",
        "glossary":          {},      # {english_term: russian_translation}
        "completed_chunks":  set(),   # индексы сохранённых фрагментов
        "chapters":          [],      # [{chunk_idx, chapter_num, title}, ...]
        "corrupted_pages":   [],      # номера страниц с плохим текстом
        "_project_loaded":   False,   # флаг: проект уже загружался?
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Загружаем project.json один раз при первом запуске сессии
    if not st.session_state._project_loaded:
        st.session_state._project_loaded = True
        load_project()



# ──────────────────────────────────────────────
#  KEYBOARD SHORTCUTS
# ──────────────────────────────────────────────

def render_hotkeys_js(prompt_text: str) -> None:
    """
    Вставляет невидимый JS-компонент, который слушает нажатия клавиш
    в родительском окне Streamlit и кликает нужные кнопки.

    Логика безопасности Ctrl+C:
      - если в окне выделен текст — работает стандартное копирование
      - если выделения нет — копируем prompt_text в буфер
    """
    # Экранируем prompt для вставки в JS-строку
    safe_prompt = (
        prompt_text
        .replace("\\", "\\\\")
        .replace("`", "\\`")
        .replace("$", "\\$")
    )

    st.components.v1.html(
        f"""
        <script>
        (function() {{
            // Предотвращаем повторную регистрацию при rerun
            if (window.parent.__hotkeys_registered) return;
            window.parent.__hotkeys_registered = true;

            const PROMPT = `{safe_prompt}`;

            // Находим кнопку по точному тексту в родительском документе
            function clickBtn(text) {{
                const doc = window.parent.document;
                const btn = [...doc.querySelectorAll('button')]
                    .find(b => b.innerText.trim() === text);
                if (btn) btn.click();
            }}

            // Копируем текст в буфер
            function copyText(text) {{
                window.parent.navigator.clipboard
                    .writeText(text)
                    .catch(() => {{/* тихая ошибка */}});
            }}

            window.parent.document.addEventListener('keydown', function(e) {{
                const isMac = /Mac|iPhone|iPad/i.test(navigator.platform);
                const ctrl  = isMac ? e.metaKey : e.ctrlKey;
                if (!ctrl) return;

                switch (e.key) {{

                    // ── Ctrl/Cmd + → ─────────────────────────────
                    case 'ArrowRight':
                        e.preventDefault();
                        clickBtn('Вперёд ▶');
                        break;

                    // ── Ctrl/Cmd + ← ─────────────────────────────
                    case 'ArrowLeft':
                        e.preventDefault();
                        clickBtn('◀ Назад');
                        break;

                    // ── Ctrl/Cmd + S → сохранить ─────────────────
                    case 's':
                    case 'S':
                        e.preventDefault();
                        clickBtn('💾 Сохранить перевод');
                        break;

                    // ── Ctrl/Cmd + Enter → сохранить ─────────────
                    case 'Enter':
                        e.preventDefault();
                        clickBtn('💾 Сохранить перевод');
                        break;

                    // ── Ctrl/Cmd + C → копировать prompt ─────────
                    // Срабатывает только когда нет выделения текста
                    case 'c':
                    case 'C': {{
                        const sel = window.parent.getSelection().toString();
                        if (!sel) {{
                            e.preventDefault();
                            copyText(PROMPT);
                        }}
                        // Есть выделение → стандартное копирование
                        break;
                    }}
                }}
            }});
        }})();
        </script>
        """,
        height=0,
    )


# ──────────────────────────────────────────────
#  ГЛАВНЫЙ ИНТЕРФЕЙС
# ──────────────────────────────────────────────

def main():
    st.set_page_config(
        page_title="PDF Translator",
        page_icon="📄",
        layout="wide",
    )
    init_state()

    st.title("📄 PDF Translator")
    st.caption("Загрузи PDF, скопируй промпт в ChatGPT/Claude, вставь перевод и сохрани результат.")

    # ── Загрузка файла ──────────────────────────
    uploaded = st.file_uploader("Загрузить PDF", type="pdf")

    if uploaded and uploaded.name != st.session_state.pdf_name:
        with st.spinner("Извлекаем текст…"):
            text, corrupted = extract_text_from_pdf(uploaded)
            st.session_state.chunks           = split_into_chunks(text, words_per_chunk=1000, overlap=150)
            st.session_state.translations     = {}
            st.session_state.active_chunk     = 0
            st.session_state.completed_chunks = set()
            st.session_state.corrupted_pages  = corrupted
            st.session_state.chapters         = detect_chapters(st.session_state.chunks)
            st.session_state.pdf_name         = uploaded.name
        ch_count = len(st.session_state.chapters)
        ch_msg   = f" · Найдено глав: **{ch_count}**" if ch_count else ""
        st.success(f"Готово! Фрагментов: {len(st.session_state.chunks)}{ch_msg}")
        if corrupted:
            st.warning(
                f"⚠️ Страниц с плохим текстом: **{len(corrupted)}** "
                f"(стр. {', '.join(str(p) for p in corrupted[:10])}"
                f"{'…' if len(corrupted) > 10 else ''}). "
                f"Помечены как `[UNREADABLE TEXT]`."
            )

    # Ничего не показываем, пока PDF не загружен
    if not st.session_state.chunks:
        st.info("👆 Загрузи PDF, чтобы начать работу.")
        return

    chunks      = st.session_state.chunks
    total       = len(chunks)
    idx         = st.session_state.active_chunk   # текущий индекс (0-based)

    # ── Боковая панель со списком фрагментов ────
    with st.sidebar:
        st.header("📑 Навигация")
        st.caption(f"Файл: **{st.session_state.pdf_name}**")
        st.caption(f"Фрагментов: **{total}**")

        # ── Список глав ──────────────────────────────
        chapters     = st.session_state.chapters
        current_ch   = chunk_to_chapter(idx, chapters)  # глава текущего чанка

        if chapters:
            st.divider()
            st.caption(f"📚 Главы ({len(chapters)}):")
            for ch in chapters:
                is_active_ch = (
                    current_ch is not None
                    and current_ch["chunk_idx"] == ch["chunk_idx"]
                )
                # Обрезаем длинные заголовки до 28 символов
                short_title = ch["title"]
                if len(short_title) > 28:
                    short_title = short_title[:25] + "…"
                ch_label = f"Гл. {ch['chapter_num']} · {short_title}"
                if st.button(ch_label, key=f"ch_btn_{ch['chunk_idx']}",
                             use_container_width=True):
                    st.session_state.active_chunk = ch["chunk_idx"]
                    st.rerun()

        st.divider()
        st.caption("Фрагменты:")

        # Кнопки фрагментов — без emoji, только текст.
        # Цвет каждой кнопки выставляется через JS-пейнтер ниже.
        for i in range(total):
            label = f"Фрагмент {i + 1}"
            if st.button(label, key=f"btn_{i}", use_container_width=True):
                st.session_state.active_chunk = i
                st.rerun()

        # Цвета для кнопок фрагментов
        btn_colors = []
        for i in range(total):
            is_done    = bool(st.session_state.translations.get(i, "").strip())
            is_current = (i == idx)
            if is_current:
                btn_colors.append({"bg": "#fff3cd", "fg": "#856404",
                                   "border": "#ffc107", "fw": "700"})
            elif is_done:
                btn_colors.append({"bg": "#d1e7dd", "fg": "#0f5132",
                                   "border": "#a3cfbb", "fw": "400"})
            else:
                btn_colors.append({"bg": "#f0f2f6", "fg": "#6c757d",
                                   "border": "#dee2e6", "fw": "400"})

        # Цвета для кнопок глав:
        #   текущая глава → жёлтый
        #   остальные     → синий/индиго
        ch_colors = []
        for ch in chapters:
            is_active_ch = (
                current_ch is not None
                and current_ch["chunk_idx"] == ch["chunk_idx"]
            )
            if is_active_ch:
                ch_colors.append({"bg": "#fff3cd", "fg": "#856404",
                                  "border": "#ffc107", "fw": "700"})
            else:
                ch_colors.append({"bg": "#e8eaf6", "fg": "#3949ab",
                                  "border": "#c5cae9", "fw": "400"})

        st.components.v1.html(
            f"""
            <script>
            const COLORS    = {json.dumps(btn_colors)};
            const CH_COLORS = {json.dumps(ch_colors)};

            function paint() {{
                const sidebar = window.parent.document
                    .querySelector('[data-testid="stSidebar"]');
                if (!sidebar) return;

                // ── Кнопки фрагментов: текст строго «Фрагмент N» ──
                const btns = [...sidebar.querySelectorAll('button')]
                    .filter(b => /^Фрагмент \\d+$/.test(b.innerText.trim()));
                btns.forEach((btn, i) => {{
                    if (i >= COLORS.length) return;
                    const c = COLORS[i];
                    btn.style.backgroundColor = c.bg;
                    btn.style.color           = c.fg;
                    btn.style.fontWeight      = c.fw;
                    btn.style.border          = '1px solid ' + c.border;
                    btn.style.borderRadius    = '6px';
                    btn.style.transition      = 'background-color 0.15s';
                }});

                // ── Кнопки глав: текст начинается с «Гл. N» ──
                const chBtns = [...sidebar.querySelectorAll('button')]
                    .filter(b => /^Гл\\. \\d+/.test(b.innerText.trim()));
                chBtns.forEach((btn, i) => {{
                    if (i >= CH_COLORS.length) return;
                    const c = CH_COLORS[i];
                    btn.style.backgroundColor = c.bg;
                    btn.style.color           = c.fg;
                    btn.style.fontWeight      = c.fw;
                    btn.style.border          = '1px solid ' + c.border;
                    btn.style.borderRadius    = '6px';
                    btn.style.transition      = 'background-color 0.15s';
                }});
            }}

            paint();
            setTimeout(paint, 150);
            setTimeout(paint, 500);
            </script>
            """,
            height=0,
        )

        # ── Прогресс ────────────────────────────────
        st.divider()
        done_count = sum(
            1 for i in range(total)
            if st.session_state.translations.get(i, "").strip()
        )
        pct = done_count / total * 100
        st.progress(done_count / total)
        st.markdown(
            f"**Переведено:** {done_count} / {total} фрагментов &nbsp;·&nbsp; "
            f"**{pct:.1f}%** книги",
            unsafe_allow_html=True,
        )

        # Статус autosave
        if PROJECT_FILE.exists():
            import time
            mtime = PROJECT_FILE.stat().st_mtime
            saved_at = time.strftime("%H:%M:%S", time.localtime(mtime))
            st.caption(f"💾 Autosave: **{saved_at}**")

        # Предупреждение о повреждённых страницах
        if st.session_state.corrupted_pages:
            cp = st.session_state.corrupted_pages
            st.warning(
                f"⚠️ **{len(cp)}** стр. с плохим текстом: "
                + ", ".join(str(p) for p in cp[:8])
                + ("…" if len(cp) > 8 else "")
            )

        # ── Keyboard Shortcuts ───────────────────────
        with st.expander("⌨️ Горячие клавиши"):
            is_mac = True   # показываем оба варианта
            st.markdown(
                """
| Действие | Mac | Windows |
|---|---|---|
| Сохранить | `⌘S` | `Ctrl+S` |
| Сохранить | `⌘↵` | `Ctrl+↵` |
| Следующий | `⌘→` | `Ctrl+→` |
| Предыдущий | `⌘←` | `Ctrl+←` |
| Копировать промпт | `⌘C`¹ | `Ctrl+C`¹ |

<small>¹ Только когда нет выделения текста</small>
                """,
                unsafe_allow_html=True,
            )

        # ── Glossary Manager ────────────────────────
        st.divider()
        gls = st.session_state.glossary          # удобный псевдоним
        term_count = len(gls)
        st.subheader(f"📖 Глоссарий ({term_count})")

        # ── Форма добавления нового термина ─────────
        st.caption("Добавить термин:")
        col_en, col_ru = st.columns(2)
        with col_en:
            new_en = st.text_input(
                "English", key="gls_en",
                placeholder="neural network",
                label_visibility="collapsed",
            )
        with col_ru:
            new_ru = st.text_input(
                "Русский", key="gls_ru",
                placeholder="нейронная сеть",
                label_visibility="collapsed",
            )
        if st.button("➕ Добавить", use_container_width=True, key="gls_add"):
            en_clean = new_en.strip()
            ru_clean = new_ru.strip()
            if en_clean and ru_clean:
                st.session_state.glossary[en_clean] = ru_clean
                save_project()
                st.rerun()
            else:
                st.warning("Заполни оба поля")

        # ── Список терминов с кнопкой удаления ──────
        if gls:
            st.caption("Термины в глоссарии:")
            for i, (en_term, ru_term) in enumerate(list(gls.items())):
                col_term, col_del = st.columns([5, 1])
                with col_term:
                    st.markdown(
                        f"<small><b>{en_term}</b> → {ru_term}</small>",
                        unsafe_allow_html=True,
                    )
                with col_del:
                    if st.button("✕", key=f"gls_del_{i}", help=f"Удалить «{en_term}»"):
                        del st.session_state.glossary[en_term]
                        save_project()
                        st.rerun()
        else:
            st.caption("_Глоссарий пуст_")

        # ── Массовый импорт (свёрнут по умолчанию) ──
        with st.expander("📋 Импорт из текста"):
            st.caption("Формат: `English term → Русский перевод`")
            bulk_raw = st.text_area(
                "bulk", height=120,
                placeholder="machine learning → машинное обучение\ndeep learning → глубокое обучение",
                label_visibility="collapsed",
                key="gls_bulk",
            )
            if st.button("Импортировать", use_container_width=True, key="gls_bulk_btn"):
                parsed = parse_glossary(bulk_raw)
                if parsed:
                    st.session_state.glossary.update(parsed)
                    save_project()
                    st.rerun()
                else:
                    st.warning("Не найдено ни одного термина")


    # ── Основная зона ────────────────────────────
    col_nav1, col_title, col_nav2 = st.columns([1, 6, 1])
    with col_nav1:
        if st.button("◀ Назад", disabled=(idx == 0)):
            st.session_state.active_chunk -= 1
            st.rerun()
    with col_title:
        ch_badge = ""
        if current_ch:
            ch_badge = f" · Гл. {current_ch['chapter_num']}: _{current_ch['title']}_"
        st.subheader(f"Фрагмент {idx + 1} из {total}{ch_badge}")
    with col_nav2:
        if st.button("Вперёд ▶", disabled=(idx == total - 1)):
            st.session_state.active_chunk += 1
            st.rerun()

    # Оригинальный текст
    st.text_area(
        "📖 Оригинальный текст",
        value=chunks[idx],
        height=200,
        disabled=True,
        key=f"orig_{idx}",
    )

    # Промпт для копирования
    prompt_text = build_prompt(chunks[idx], idx + 1, st.session_state.glossary)

    st.text_area(
        "📋 Промпт (скопируй в ChatGPT / Claude)",
        value=prompt_text,
        height=180,
        key=f"prompt_{idx}",
    )

    # Кнопка копирования (через JS-трюк Streamlit)
    escaped = prompt_text.replace("`", "\\`").replace("$", "\\$")
    st.components.v1.html(
        f"""
        <button onclick="navigator.clipboard.writeText(`{escaped}`)
                          .then(() => this.innerText = '✅ Скопировано!')
                          .catch(() => this.innerText = '❌ Ошибка')"
                style="padding:6px 14px;border-radius:6px;border:1px solid #aaa;
                       cursor:pointer;font-size:14px;background:#f0f2f6">
          📋 Копировать промпт
        </button>
        """,
        height=50,
    )

    # Поле для вставки перевода
    st.divider()
    saved_translation = st.session_state.translations.get(idx, "")
    new_translation = st.text_area(
        "✍️ Вставь перевод сюда",
        value=saved_translation,
        height=220,
        placeholder="Вставь ответ ChatGPT / Claude здесь…",
        key=f"trans_input_{idx}",
    )

    col_save, col_clear = st.columns([2, 1])
    with col_save:
        if st.button("💾 Сохранить перевод", type="primary", use_container_width=True):
            st.session_state.translations[idx] = new_translation
            st.session_state.completed_chunks.add(idx)
            save_project()          # ← autosave
            st.success("Перевод сохранён!")
            # Автоматически переходим к следующему фрагменту
            if idx < total - 1:
                st.session_state.active_chunk += 1
                st.rerun()
    with col_clear:
        if st.button("🗑 Очистить", use_container_width=True):
            st.session_state.translations[idx] = ""
            st.session_state.completed_chunks.discard(idx)
            save_project()          # ← autosave
            st.rerun()

    # ── Экспорт ──────────────────────────────────
    st.divider()
    st.subheader("💾 Экспорт")
    col_json, col_docx = st.columns(2)

    with col_json:
        json_str = save_translations_to_json(chunks, st.session_state.translations)
        st.download_button(
            label="📥 Скачать JSON",
            data=json_str.encode("utf-8"),
            file_name="translations.json",
            mime="application/json",
            use_container_width=True,
        )

    with col_docx:
        book_title  = Path(st.session_state.pdf_name).stem if st.session_state.pdf_name else "Перевод книги"
        docx_buffer = export_to_docx(chunks, st.session_state.translations,
                                     st.session_state.chapters, book_title)
        st.download_button(
            label="📥 Скачать DOCX",
            data=docx_buffer,
            file_name="translations.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    # ── Горячие клавиши (невидимый JS) ───────────
    render_hotkeys_js(prompt_text)


if __name__ == "__main__":
    main()
